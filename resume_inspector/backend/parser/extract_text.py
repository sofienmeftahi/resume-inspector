import fitz  # PyMuPDF
from docx import Document
import os
import re

def extract_text_from_file(file_path):
    """
    Extract text from PDF or DOCX files
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                return "\n".join([page.extract_text() or "" for page in pdf.pages])
        except Exception as e:
            print(f"PDF extraction error: {e}")
            return ""
    elif ext == ".docx":
        try:
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""
    else:
        return ""

def extract_text_from_pdf(file_path):
    """
    Extract text from PDF using PyMuPDF
    """
    try:
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        
        doc.close()
        return clean_text(text)
    except Exception as e:
        print(f"Error reading PDF {file_path}: {str(e)}")
        return ""

def extract_text_from_docx(file_path):
    """
    Extract text from DOCX using python-docx
    """
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return clean_text(text)
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {str(e)}")
        return ""

def clean_text(text):
    """
    Clean and normalize extracted text
    """
    if not text:
        return ""
    
    # Remove extra whitespace and normalize
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    text = text.strip()
    
    # Remove common PDF artifacts
    text = re.sub(r'[^\w\s\-\.\,\;\:\!\?\(\)\[\]\{\}\@\#\$\%\&\*\+\=\|\/\\]', '', text)
    
    return text

def get_file_info(file_path):
    """
    Get basic information about the file
    """
    try:
        file_size = os.path.getsize(file_path)
        file_extension = file_path.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            doc = fitz.open(file_path)
            page_count = len(doc)
            doc.close()
        elif file_extension in ['doc', 'docx']:
            doc = Document(file_path)
            page_count = len(doc.paragraphs) // 20  # Rough estimate
        else:
            page_count = 0
        
        return {
            'size': file_size,
            'extension': file_extension,
            'page_count': page_count
        }
    except Exception as e:
        print(f"Error getting file info: {str(e)}")
        return {'size': 0, 'extension': 'unknown', 'page_count': 0} 